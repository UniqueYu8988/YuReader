import io
import json
import sys
import tempfile
import unittest
from datetime import date
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import app  # noqa: E402
from yureader.oral_focus import SUBJECT_SOURCES, build_dataset, split_bilingual_definition_title  # noqa: E402


class OralFocusHandler(app.ReaderHandler):
    def __init__(self, payload):
        encoded = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.path = "/api/oral-focus/progress"
        self.headers = {"Content-Length": str(len(encoded))}
        self.rfile = io.BytesIO(encoded)
        self.wfile = io.BytesIO()
        self.status = None

    def send_response(self, code, message=None):
        self.status = code

    def send_header(self, key, value):
        pass

    def end_headers(self):
        pass

    def log_message(self, format, *args):
        pass


class OralFocusImportTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        for subject in SUBJECT_SOURCES:
            for item_type, filenames in subject["files"].items():
                for file_index, filename in enumerate(filenames):
                    document = Document()
                    if file_index == 0:
                        document.add_paragraph("一、测试章节\t名词解释" if item_type == "definition" else "一、测试章节\t简答论述")
                        if item_type == "definition":
                            document.add_paragraph("1.test term：测试术语")
                            document.add_paragraph("（1）这是定义。微信搜索银河研旅公众号，记乎app搜索途中口腔医学考研2班")
                            document.add_paragraph("答案末尾未正确换段 2.inline term：行内术语★★")
                            document.add_paragraph("（1）这是行内题目的答案。")
                            document.add_paragraph("3.next term：下一术语★")
                            document.add_paragraph("（1）这是跳号后的独立定义。")
                        else:
                            document.add_paragraph("1．测试论述题★★")
                            document.add_paragraph("（1）第一评分点")
                            table = document.add_table(rows=2, cols=2)
                            table.cell(0, 0).text = "项目"
                            table.cell(0, 1).text = "内容"
                            table.cell(1, 0).text = "A"
                            table.cell(1, 1).text = "B"
                    else:
                        document.add_paragraph("（2）这是上一卷的续文")
                        document.add_paragraph("二、下一章节")
                        document.add_paragraph("1．第二道论述题★")
                        document.add_paragraph("（1）第二题答案")
                    document.save(self.root / filename)

    def tearDown(self):
        self.temp.cleanup()

    def test_import_keeps_tables_joins_split_files_and_removes_promotions(self):
        payload = build_dataset(self.root)
        self.assertEqual(payload["summary"]["subject_count"], 5)
        self.assertEqual(payload["summary"]["definition_count"], 15)
        self.assertEqual(payload["summary"]["essay_count"], 7)
        self.assertEqual(payload["summary"]["table_item_count"], 5)
        self.assertGreaterEqual(payload["summary"]["excluded_promotion_blocks"], 5)
        oral_surgery = payload["subjects"][0]
        definitions = [item for chapter in oral_surgery["chapters"] for item in chapter["items"] if item["type"] == "definition"]
        first_definition, inline_definition, skipped_number_definition = definitions
        self.assertEqual(first_definition["title"], "test term")
        self.assertEqual(first_definition["definition_translation"], "测试术语")
        self.assertEqual(first_definition["source_title"], "test term：测试术语")
        self.assertEqual(first_definition["answer_markdown"], "（1）这是定义。\n\n答案末尾未正确换段")
        self.assertEqual(inline_definition["title"], "inline term")
        self.assertEqual(inline_definition["definition_translation"], "行内术语")
        self.assertEqual(inline_definition["answer_markdown"], "（1）这是行内题目的答案。")
        self.assertEqual(skipped_number_definition["title"], "next term")
        self.assertEqual(skipped_number_definition["definition_translation"], "下一术语")
        self.assertNotIn("next term", first_definition["answer_markdown"])
        self.assertEqual(payload["summary"]["bilingual_definition_title_count"], 15)
        first_essay = next(item for chapter in oral_surgery["chapters"] for item in chapter["items"] if item["type"] == "essay")
        self.assertIn("| 项目 | 内容 |", first_essay["answer_markdown"])
        self.assertIn("上一卷的续文", first_essay["answer_markdown"])
        self.assertNotIn("银河研旅", json.dumps(payload, ensure_ascii=False))

    def test_bilingual_definition_title_keeps_the_english_recall_prompt(self):
        self.assertEqual(split_bilingual_definition_title("anesthesia：麻醉"), ("anesthesia", "麻醉"))
        self.assertEqual(split_bilingual_definition_title("Vesicle/blister：疱"), ("Vesicle/blister", "疱"))
        self.assertEqual(split_bilingual_definition_title("Ante's law: Ante 法则"), ("Ante's law", "Ante 法则"))
        self.assertEqual(split_bilingual_definition_title("牙根拔除术（exodontia）"), ("牙根拔除术（exodontia）", ""))
        self.assertEqual(split_bilingual_definition_title("guided tissue regeneration, GTR"), ("guided tissue regeneration, GTR", ""))


class OralFocusRuntimeTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.data = self.root / "data"
        self.focus = self.data / "oral-focus"
        item_id = "oral-focus-0123456789abcdef"
        self.item_id = item_id
        dataset = {
            "schema_version": 1,
            "summary": {"subject_count": 1, "item_count": 1},
            "subjects": [
                {
                    "id": "oral-surgery",
                    "short_title": "口外",
                    "title": "口腔颌面外科学",
                    "book_id": "oral-maxillofacial-surgery-8e",
                    "item_count": 1,
                    "chapter_count": 1,
                    "chapters": [
                        {
                            "id": "oral-surgery-ch02",
                            "order": 2,
                            "title": "二、基础知识",
                            "definition_count": 0,
                            "essay_count": 1,
                            "starred_count": 1,
                            "items": [
                                {
                                    "id": item_id,
                                    "order": 1,
                                    "type": "essay",
                                    "type_label": "简答论述",
                                    "title": "活组织检查的注意事项",
                                    "star_level": 3,
                                    "answer_markdown": "（1）正确取材\n\n（2）结合临床",
                                    "character_count": 18,
                                    "has_table": False,
                                    "has_unreviewed_image": False,
                                    "source_files": ["口外论述1.docx"],
                                    "source_paragraph": 10,
                                }
                            ],
                        }
                    ],
                }
            ],
        }
        self.focus.mkdir(parents=True)
        (self.focus / "content.json").write_text(json.dumps(dataset, ensure_ascii=False), encoding="utf-8")
        names = ("DATA_DIR", "ORAL_FOCUS_DIR", "ORAL_FOCUS_CONTENT_PATH", "ORAL_FOCUS_PROGRESS_PATH", "ACTIVITY_PATH", "ORAL_FOCUS_CACHE", "obsidian_vault")
        self.originals = {name: getattr(app, name) for name in names}
        app.DATA_DIR = self.data
        app.ORAL_FOCUS_DIR = self.focus
        app.ORAL_FOCUS_CONTENT_PATH = self.focus / "content.json"
        app.ORAL_FOCUS_PROGRESS_PATH = self.focus / "progress.json"
        app.ACTIVITY_PATH = self.data / "activity.json"
        app.ORAL_FOCUS_CACHE = {"mtime_ns": None, "payload": None, "items": {}}
        app.obsidian_vault = lambda: None

    def tearDown(self):
        for name, value in self.originals.items():
            setattr(app, name, value)
        self.temp.cleanup()

    def test_reference_is_absent_until_explicit_reveal(self):
        index = app.oral_focus_index_payload()
        self.assertTrue(index["available"])
        self.assertNotIn("answer_markdown", json.dumps(index, ensure_ascii=False))
        self.assertFalse(index["subjects"][0]["chapters"][0]["items"][0]["completed"])
        hidden = app.oral_focus_item_payload(self.item_id)
        self.assertNotIn("answer_markdown", hidden)
        self.assertEqual(hidden["storage"], "local")
        self.assertEqual(hidden["obsidian_uri"], "obsidian://open")
        revealed = app.oral_focus_item_payload(self.item_id, reveal=True)
        self.assertIn("正确取材", revealed["answer_markdown"])

    def test_chapter_payload_keeps_answers_out_until_global_reveal(self):
        hidden = app.oral_focus_chapter_payload("oral-surgery", "oral-surgery-ch02", "essay")
        self.assertEqual(len(hidden["items"]), 1)
        self.assertNotIn("answer_markdown", hidden["items"][0])
        revealed = app.oral_focus_chapter_payload("oral-surgery", "oral-surgery-ch02", "essay", reveal=True)
        self.assertIn("正确取材", revealed["items"][0]["answer_markdown"])

    def test_progress_enters_subjective_activity_and_review(self):
        handler = OralFocusHandler({"item_id": self.item_id, "answer": "我的闭卷答案", "memory_note": "漏了结合临床", "mastery": "fuzzy"})
        handler.do_POST()
        self.assertEqual(handler.status, 200)
        response = json.loads(handler.wfile.getvalue().decode("utf-8"))
        self.assertTrue(response["saved"])
        self.assertTrue(app.ORAL_FOCUS_PROGRESS_PATH.is_file())
        self.assertIn("漏了结合临床", Path(response["path"]).read_text(encoding="utf-8"))
        self.assertTrue(app.oral_focus_index_payload()["subjects"][0]["chapters"][0]["items"][0]["completed"])
        activities = app.activity_records_payload(date.today().isoformat(), "subjective_practice")["activities"]
        self.assertEqual(len(activities), 1)
        self.assertEqual(activities[0]["resume_target"]["view"], "oral_focus")
        sources = app.review_source_records(date.today().isoformat(), [], {})
        self.assertEqual(len(sources), 1)
        self.assertEqual(sources[0]["title"], "活组织检查的注意事项")
        self.assertIn("我的闭卷答案", sources[0]["markdown"])
        self.assertIn("漏了结合临床", sources[0]["markdown"])

    def test_sidebar_note_preserves_hidden_historical_fields(self):
        first = app.save_oral_focus_progress(self.item_id, "历史作答", "旧笔记", "mastered")
        second = app.save_oral_focus_progress(
            self.item_id,
            first["progress"]["answer"],
            "从 Gemini 整理的新笔记",
            first["progress"]["mastery"],
        )
        self.assertEqual(second["progress"]["answer"], "历史作答")
        self.assertEqual(second["progress"]["mastery"], "mastered")
        note = Path(second["path"]).read_text(encoding="utf-8")
        self.assertIn("### 学习笔记", note)
        self.assertIn("从 Gemini 整理的新笔记", note)


if __name__ == "__main__":
    unittest.main()
