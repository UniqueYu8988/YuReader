"""Build a local, source-preserving oral-medicine focus dataset from DOCX files."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path


SCHEMA_VERSION = 1
SUBJECT_SOURCES = [
    {
        "id": "oral-surgery",
        "short_title": "口外",
        "title": "口腔颌面外科学",
        "book_id": "oral-maxillofacial-surgery-8e",
        "files": {"definition": ["口外名解.docx"], "essay": ["口外论述1.docx", "口外论述2.docx"]},
    },
    {
        "id": "oral-pathology",
        "short_title": "口组",
        "title": "口腔组织病理学",
        "book_id": "oral-pathology-8e",
        "files": {"definition": ["口组名解.docx"], "essay": ["口组论述.docx"]},
    },
    {
        "id": "dental-pulp",
        "short_title": "牙体",
        "title": "牙体牙髓病学",
        "book_id": "dental-pulp-5e",
        "files": {"definition": ["牙体名解.docx"], "essay": ["牙体论述.docx"]},
    },
    {
        "id": "periodontology",
        "short_title": "牙周",
        "title": "牙周病学",
        "book_id": "periodontology-5e",
        "files": {"definition": ["牙周名解.docx"], "essay": ["牙周论述.docx"]},
    },
    {
        "id": "prosthodontics",
        "short_title": "修复",
        "title": "口腔修复学",
        "book_id": "prosthodontics-8e",
        "files": {"definition": ["修复名解.docx"], "essay": ["修复论述1.docx", "修复论述2.docx"]},
    },
]
TYPE_LABELS = {"definition": "名词解释", "essay": "简答论述"}
CHINESE_DIGITS = {"零": 0, "〇": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
CHAPTER_RE = re.compile(r"^\s*([一二三四五六七八九十百〇零]+)\s*[、.．]\s*(.+?)\s*$")
QUESTION_RE = re.compile(r"^\s*(\d{1,3})\s*[.．、]\s*(.+?)\s*$")
EMBEDDED_DEFINITION_QUESTION_RE = re.compile(r"(?m)^\s*\d{1,3}\s*[.．、]\s*(?!\d)\S+")
INLINE_DEFINITION_QUESTION_RE = re.compile(
    r"(?<![\dA-Za-z])\d{1,3}\s*[.．、]\s*"
    r"[A-Za-z][^：:\n]{0,140}[：:]\s*[\u3400-\u9fff][^\n]*$"
)
PROMOTION_RE = re.compile(
    r"(?:微信搜索(?:公众号[：:]?)?\s*银河研旅(?:公众号)?(?:，|,)?\s*)?"
    r"(?:记乎\s*app\s*搜索(?:班级[：:]?)?\s*途中口腔医学考研(?:2班)?)",
    re.IGNORECASE,
)
PURCHASER_RE = re.compile(r"\s*购买者[^\n，。；]{0,40}(?:大学|学院)?\s*$")


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _stable_id(*parts: object) -> str:
    raw = "\u241f".join(str(part or "").strip() for part in parts)
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]


def chinese_number(value: str) -> int:
    """Parse the chapter numerals used by these documents (1-99)."""
    value = str(value or "").strip()
    if not value:
        return 0
    if "十" in value:
        left, right = value.split("十", 1)
        tens = CHINESE_DIGITS.get(left, 1) if left else 1
        ones = CHINESE_DIGITS.get(right, 0) if right else 0
        return tens * 10 + ones
    return CHINESE_DIGITS.get(value, 0)


def clean_source_text(value: object) -> tuple[str, bool]:
    text = str(value or "").replace("\u00a0", " ").replace("\r\n", "\n").strip()
    cleaned = PROMOTION_RE.sub("", text)
    cleaned = PURCHASER_RE.sub("", cleaned)
    cleaned = re.sub(r"\s+([，。；：])", r"\1", cleaned).strip(" \t，,")
    return cleaned, cleaned != text


def clean_chapter_title(raw: str) -> str:
    title = str(raw or "").split("\t", 1)[0].strip()
    title = re.sub(r"\s*(?:组病)?(?:名词解释|名解|简答论述|论述)\s*$", "", title).strip()
    return title or "未分章"


def clean_question_title(raw: str) -> tuple[str, int]:
    stars = min(3, str(raw or "").count("★") + str(raw or "").count("＊"))
    title = re.sub(r"[★＊*]+", "", str(raw or ""))
    title, _ = clean_source_text(title)
    return re.sub(r"\s+", " ", title).strip(" ：:"), stars


def split_embedded_definition(title: str) -> tuple[str, str]:
    """Separate definitions that OCR kept on the same line as the term."""
    value = str(title or "").strip()
    if "。" in value:
        lead, remainder = value.split("。", 1)
        if remainder.strip():
            return lead.strip(), remainder.strip()
    if "：" in value:
        lead, remainder = value.split("：", 1)
        if len(remainder) >= 18 or re.search(r"(?:指|是|也称|利用|若|由|有|称为)", remainder):
            return lead.strip(), remainder.strip()
    return value, ""


def split_bilingual_definition_title(title: str) -> tuple[str, str]:
    """Keep the English recall prompt and separate its Chinese translation."""
    value = str(title or "").strip()
    match = re.match(r"^([A-Za-z][^：:]*?)\s*[：:]\s*(.*[\u3400-\u9fff].*)$", value)
    if not match:
        return value, ""
    english_prompt = re.sub(r"\s+", " ", match.group(1)).strip()
    chinese_translation = re.sub(r"\s+", " ", match.group(2)).strip(" ：:")
    return english_prompt or value, chinese_translation


def table_markdown(table: object) -> str:
    rows: list[list[str]] = []
    for row in getattr(table, "rows", []):
        values = []
        for cell in row.cells:
            value, _ = clean_source_text(cell.text)
            values.append(value.replace("|", "\\|").replace("\n", "<br>"))
        if any(values):
            rows.append(values)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    return "\n".join(
        ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
        + ["| " + " | ".join(row) + " |" for row in rows[1:]]
    )


def iter_doc_blocks(document: object):
    """Yield paragraphs and tables in document order without rewriting OOXML."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def build_dataset(source_dir: Path | str) -> dict:
    """Read the known OCR DOCX set and return one deterministic learning dataset."""
    from docx import Document

    root = Path(source_dir).resolve()
    if not root.is_dir():
        raise ValueError("source directory not found")
    subjects: list[dict] = []
    source_files: list[dict] = []
    warnings: list[dict] = []
    excluded_promotions = 0
    image_markers = 0
    bilingual_definition_titles = 0

    for subject_config in SUBJECT_SOURCES:
        chapter_map: dict[int, dict] = {}
        chapter_order: list[int] = []
        subject_item_count = 0
        for item_type in ("definition", "essay"):
            current_chapter: dict | None = None
            current_item: dict | None = None
            expected_question = 1
            duplicate_titles: dict[tuple[int, str], int] = {}

            def finish_item() -> None:
                nonlocal current_item, subject_item_count
                if not current_item:
                    return
                answer = "\n\n".join(block for block in current_item.pop("_blocks", []) if block.strip()).strip()
                current_item["answer_markdown"] = answer or "暂无可识别的参考答案。"
                current_item["character_count"] = len(answer)
                current_item["has_table"] = bool(current_item.pop("_has_table", False))
                current_item["has_unreviewed_image"] = bool(current_item.pop("_has_image", False))
                current_item["source_files"] = list(dict.fromkeys(current_item.get("source_files") or []))
                current_item["order"] = len(current_chapter["items"]) + 1
                if item_type == "definition" and (
                    EMBEDDED_DEFINITION_QUESTION_RE.search(answer)
                    or INLINE_DEFINITION_QUESTION_RE.search(answer)
                ):
                    raise ValueError(f"definition still contains a nested question boundary: {current_item['id']}")
                current_chapter["items"].append(current_item)
                subject_item_count += 1
                if current_item["character_count"] > 5000:
                    warnings.append({"kind": "oversized_item", "item_id": current_item["id"], "characters": current_item["character_count"]})
                elif item_type == "definition" and current_item["character_count"] > 2500:
                    warnings.append({"kind": "oversized_definition", "item_id": current_item["id"], "characters": current_item["character_count"]})
                current_item = None

            for filename in subject_config["files"][item_type]:
                source = (root / filename).resolve()
                if source.parent != root or not source.is_file():
                    raise ValueError(f"required source missing: {filename}")
                document = Document(source)
                file_sha = _sha256(source)
                source_files.append(
                    {
                        "name": filename,
                        "sha256": file_sha,
                        "bytes": source.stat().st_size,
                        "paragraphs": len(document.paragraphs),
                        "tables": len(document.tables),
                        "inline_shapes": len(document.inline_shapes),
                    }
                )
                paragraph_number = 0
                for block_kind, block in iter_doc_blocks(document):
                    if block_kind == "paragraph":
                        paragraph_number += 1
                        text, removed = clean_source_text(block.text)
                        excluded_promotions += int(removed)
                        chapter_match = CHAPTER_RE.match(text)
                        if chapter_match:
                            finish_item()
                            number = chinese_number(chapter_match.group(1)) or (len(chapter_order) + 1)
                            if number not in chapter_map:
                                chapter_order.append(number)
                                chapter_map[number] = {
                                    "id": f"{subject_config['id']}-ch{number:02d}",
                                    "order": number,
                                    "title": clean_chapter_title(text),
                                    "items": [],
                                }
                            current_chapter = chapter_map[number]
                            expected_question = 1
                            continue
                        if item_type == "definition" and current_item:
                            inline_question = INLINE_DEFINITION_QUESTION_RE.search(text)
                            if inline_question and inline_question.start() > 0:
                                preceding_answer = text[: inline_question.start()].strip()
                                if preceding_answer:
                                    current_item["_blocks"].append(preceding_answer)
                                    current_item["source_files"].append(filename)
                                finish_item()
                                text = inline_question.group(0).strip()
                        question_match = QUESTION_RE.match(text)
                        question_number = int(question_match.group(1)) if question_match else 0
                        question_tail = question_match.group(2).strip() if question_match else ""
                        is_question_boundary = bool(
                            question_match
                            and not re.match(r"^\d", question_tail)
                            and (item_type == "definition" or question_number == expected_question)
                        )
                        if is_question_boundary:
                            finish_item()
                            if current_chapter is None:
                                number = 0
                                if number not in chapter_map:
                                    chapter_order.insert(0, number)
                                    chapter_map[number] = {"id": f"{subject_config['id']}-ch00", "order": 0, "title": "未分章", "items": []}
                                current_chapter = chapter_map[number]
                            title, stars = clean_question_title(question_match.group(2))
                            embedded_answer = ""
                            source_title = title
                            id_title = title
                            chinese_translation = ""
                            if item_type == "definition":
                                legacy_id_title, legacy_embedded_answer = split_embedded_definition(title)
                                id_title = legacy_id_title
                                title, chinese_translation = split_bilingual_definition_title(title)
                                if chinese_translation:
                                    bilingual_definition_titles += 1
                                else:
                                    title, embedded_answer = legacy_id_title, legacy_embedded_answer
                            duplicate_key = (int(current_chapter["order"]), id_title)
                            duplicate_titles[duplicate_key] = duplicate_titles.get(duplicate_key, 0) + 1
                            item_id = "oral-focus-" + _stable_id(subject_config["id"], item_type, current_chapter["id"], id_title, duplicate_titles[duplicate_key])
                            current_item = {
                                "id": item_id,
                                "type": item_type,
                                "type_label": TYPE_LABELS[item_type],
                                "title": title or f"第 {question_match.group(1)} 题",
                                "source_number": question_number,
                                "star_level": stars,
                                "source_files": [filename],
                                "source_paragraph": paragraph_number,
                                "source_sha256": file_sha,
                                "_blocks": [embedded_answer] if embedded_answer else [],
                                "_has_table": False,
                                "_has_image": False,
                            }
                            if chinese_translation:
                                current_item["definition_translation"] = chinese_translation
                                current_item["source_title"] = source_title
                            expected_question += 1
                            continue
                        if current_item and text:
                            current_item["_blocks"].append(text)
                            current_item["source_files"].append(filename)
                        if current_item and block._p.xpath(".//a:blip"):
                            current_item["_blocks"].append("> 原始文档此处包含图片，尚待人工确认后发布。")
                            current_item["_has_image"] = True
                            image_markers += 1
                    elif current_item:
                        markdown = table_markdown(block)
                        if markdown:
                            current_item["_blocks"].append(markdown)
                            current_item["_has_table"] = True
            finish_item()

        chapters = [chapter_map[number] for number in sorted(chapter_order) if chapter_map[number]["items"]]
        for chapter in chapters:
            chapter["definition_count"] = sum(item["type"] == "definition" for item in chapter["items"])
            chapter["essay_count"] = sum(item["type"] == "essay" for item in chapter["items"])
            chapter["starred_count"] = sum(item["star_level"] > 0 for item in chapter["items"])
        subjects.append(
            {
                **{key: subject_config[key] for key in ("id", "short_title", "title", "book_id")},
                "item_count": subject_item_count,
                "chapter_count": len(chapters),
                "chapters": chapters,
            }
        )

    return {
        "schema_version": SCHEMA_VERSION,
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "source": {"external_path": str(root), "files": source_files},
        "summary": {
            "subject_count": len(subjects),
            "chapter_count": sum(subject["chapter_count"] for subject in subjects),
            "item_count": sum(subject["item_count"] for subject in subjects),
            "definition_count": sum(item["type"] == "definition" for subject in subjects for chapter in subject["chapters"] for item in chapter["items"]),
            "essay_count": sum(item["type"] == "essay" for subject in subjects for chapter in subject["chapters"] for item in chapter["items"]),
            "table_item_count": sum(item["has_table"] for subject in subjects for chapter in subject["chapters"] for item in chapter["items"]),
            "unreviewed_image_item_count": sum(item["has_unreviewed_image"] for subject in subjects for chapter in subject["chapters"] for item in chapter["items"]),
            "excluded_promotion_blocks": excluded_promotions,
            "image_markers": image_markers,
            "bilingual_definition_title_count": bilingual_definition_titles,
        },
        "warnings": warnings,
        "subjects": subjects,
    }


def write_dataset(source_dir: Path | str, output_path: Path | str) -> dict:
    payload = build_dataset(source_dir)
    target = Path(output_path).resolve()
    target.parent.mkdir(parents=True, exist_ok=True)
    temporary = target.with_suffix(target.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(target)
    return payload


# Backend server bridge routines
from urllib.parse import quote
from yureader.config import (
    DATA_DIR,
    ORAL_FOCUS_DIR,
    ORAL_FOCUS_CONTENT_PATH,
    ORAL_FOCUS_PROGRESS_PATH,
    ORAL_FOCUS_LOCK,
    ORAL_FOCUS_CACHE,
)
from yureader.utils import safe_note_component, atomic_write
from yureader.obsidian import obsidian_vault

def load_oral_focus() -> tuple[dict, dict[str, dict]]:
    """Load the ignored local focus dataset and index items without mutating it."""
    if not ORAL_FOCUS_CONTENT_PATH.is_file():
        return {"schema_version": 1, "summary": {}, "subjects": []}, {}
    mtime_ns = ORAL_FOCUS_CONTENT_PATH.stat().st_mtime_ns
    if ORAL_FOCUS_CACHE.get("mtime_ns") == mtime_ns and isinstance(ORAL_FOCUS_CACHE.get("payload"), dict):
        return ORAL_FOCUS_CACHE["payload"], ORAL_FOCUS_CACHE.get("items", {})
    payload = json.loads(ORAL_FOCUS_CONTENT_PATH.read_text(encoding="utf-8-sig"))
    if not isinstance(payload, dict) or int(payload.get("schema_version") or 0) != 1 or not isinstance(payload.get("subjects"), list):
        raise ValueError("invalid oral focus dataset")
    items: dict[str, dict] = {}
    for subject in payload["subjects"]:
        if not isinstance(subject, dict):
            continue
        for chapter in subject.get("chapters") if isinstance(subject.get("chapters"), list) else []:
            if not isinstance(chapter, dict):
                continue
            for item in chapter.get("items") if isinstance(chapter.get("items"), list) else []:
                if not isinstance(item, dict):
                    continue
                item_id = str(item.get("id") or "")
                if not re.fullmatch(r"oral-focus-[a-f0-9]{16}", item_id) or item_id in items:
                    raise ValueError("invalid oral focus item id")
                items[item_id] = {**item, "subject": subject, "chapter": chapter}
    ORAL_FOCUS_CACHE.update({"mtime_ns": mtime_ns, "payload": payload, "items": items})
    return payload, items


def load_oral_focus_progress() -> dict:
    if not ORAL_FOCUS_PROGRESS_PATH.is_file():
        return {"schema_version": 1, "items": {}}
    try:
        payload = json.loads(ORAL_FOCUS_PROGRESS_PATH.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError):
        return {"schema_version": 1, "items": {}}
    if not isinstance(payload, dict) or not isinstance(payload.get("items"), dict):
        return {"schema_version": 1, "items": {}}
    return payload


def oral_focus_index_payload() -> dict:
    dataset, _items = load_oral_focus()
    progress = load_oral_focus_progress().get("items", {})
    subjects: list[dict] = []
    for subject in dataset.get("subjects") or []:
        chapters: list[dict] = []
        for chapter in subject.get("chapters") or []:
            public_items = []
            for item in chapter.get("items") or []:
                state = progress.get(str(item.get("id") or ""), {})
                state = state if isinstance(state, dict) else {}
                completed = bool(
                    str(state.get("answer") or "").strip()
                    or str(state.get("memory_note") or "").strip()
                    or str(state.get("mastery") or "unseen") != "unseen"
                )
                public_items.append(
                    {
                        key: item.get(key)
                        for key in ("id", "order", "type", "type_label", "title", "star_level", "character_count", "has_table", "has_unreviewed_image")
                    }
                    | {"mastery": str(state.get("mastery") or "unseen"), "completed": completed}
                )
            chapters.append(
                {
                    key: chapter.get(key)
                    for key in ("id", "order", "title", "definition_count", "essay_count", "starred_count")
                }
                | {"items": public_items}
            )
        subject_progress = [progress.get(item["id"], {}) for chapter in chapters for item in chapter["items"]]
        subjects.append(
            {
                key: subject.get(key)
                for key in ("id", "short_title", "title", "book_id", "item_count", "chapter_count")
            }
            | {
                "studied_count": sum(
                    isinstance(item, dict)
                    and bool(
                        str(item.get("answer") or "").strip()
                        or str(item.get("memory_note") or "").strip()
                        or item.get("mastery") not in (None, "", "unseen")
                    )
                    for item in subject_progress
                ),
                "mastered_count": sum(isinstance(item, dict) and item.get("mastery") == "mastered" for item in subject_progress),
                "chapters": chapters,
            }
        )
    return {"available": bool(subjects), "summary": dataset.get("summary") or {}, "subjects": subjects}


def _oral_focus_public_record(record: dict, progress: dict, *, reveal: bool = False) -> dict:
    subject = record["subject"]
    chapter = record["chapter"]
    progress = progress if isinstance(progress, dict) else {}
    _note_target, note_storage, note_uri = oral_focus_notes_target(subject)
    public = {
        key: record.get(key)
        for key in ("id", "order", "type", "type_label", "title", "star_level", "character_count", "has_table", "has_unreviewed_image", "source_files", "source_paragraph")
    }
    public.update(
        {
            "subject": {key: subject.get(key) for key in ("id", "short_title", "title", "book_id")},
            "chapter": {key: chapter.get(key) for key in ("id", "order", "title")},
            "progress": {
                "answer": str(progress.get("answer") or ""),
                "memory_note": str(progress.get("memory_note") or ""),
                "mastery": str(progress.get("mastery") or "unseen"),
                "updated_at": str(progress.get("updated_at") or ""),
            },
            "reference_revealed": bool(reveal),
            "storage": note_storage,
            "obsidian_uri": note_uri,
        }
    )
    if reveal:
        public["answer_markdown"] = str(record.get("answer_markdown") or "")
        public["definition_translation"] = str(record.get("definition_translation") or "")
    return public


def oral_focus_item_payload(item_id: str, *, reveal: bool = False) -> dict:
    if not re.fullmatch(r"oral-focus-[a-f0-9]{16}", str(item_id or "")):
        raise ValueError("invalid oral focus item id")
    _dataset, items = load_oral_focus()
    record = items.get(item_id)
    if not record:
        raise ValueError("oral focus item not found")
    progress = load_oral_focus_progress().get("items", {}).get(item_id, {})
    return _oral_focus_public_record(record, progress, reveal=reveal)


def oral_focus_chapter_payload(subject_id: str, chapter_id: str, item_type: str, *, reveal: bool = False) -> dict:
    if not re.fullmatch(r"[a-z0-9-]+", str(subject_id or "")) or not re.fullmatch(r"[a-z0-9-]+", str(chapter_id or "")):
        raise ValueError("invalid oral focus chapter")
    if item_type not in {"", "definition", "essay"}:
        raise ValueError("invalid oral focus type")
    dataset, items = load_oral_focus()
    subject = next((entry for entry in dataset.get("subjects") or [] if entry.get("id") == subject_id), None)
    if not subject:
        raise ValueError("oral focus subject not found")
    chapter = next((entry for entry in subject.get("chapters") or [] if entry.get("id") == chapter_id), None)
    if not chapter:
        raise ValueError("oral focus chapter not found")
    progress_items = load_oral_focus_progress().get("items", {})
    public_items = []
    for item in chapter.get("items") or []:
        if item_type and item.get("type") != item_type:
            continue
        record = items.get(str(item.get("id") or ""))
        if record:
            public_items.append(_oral_focus_public_record(record, progress_items.get(record["id"], {}), reveal=reveal))
    return {
        "subject": {key: subject.get(key) for key in ("id", "short_title", "title", "book_id")},
        "chapter": {key: chapter.get(key) for key in ("id", "order", "title")},
        "type": item_type,
        "reference_revealed": bool(reveal),
        "items": public_items,
    }


def oral_focus_notes_target(subject: dict) -> tuple[Path, str, str]:
    subject_name = safe_note_component(subject.get("title") or subject.get("id"), "口腔重点")
    vault = obsidian_vault()
    if vault:
        root = (vault / "YuReader" / "医学" / subject_name).resolve()
        target = (root / "重点背诵.md").resolve()
        if target.parent != root:
            raise ValueError("oral focus note path escapes Obsidian vault")
        return target, "obsidian", f"obsidian://open?path={quote(str(target.relative_to(vault)).replace(os.sep, '/'))}"
    target = (ORAL_FOCUS_DIR / "notes" / f"{subject.get('id')}.md").resolve()
    if ORAL_FOCUS_DIR.resolve() not in target.parents:
        raise ValueError("oral focus note path escapes data directory")
    return target, "local", "obsidian://open"


def write_oral_focus_notes(subject_id: str, progress_payload: dict) -> tuple[Path, str, str]:
    dataset, items = load_oral_focus()
    subject = next((entry for entry in dataset.get("subjects") or [] if entry.get("id") == subject_id), None)
    if not subject:
        raise ValueError("oral focus subject not found")
    lines = [f"# {subject.get('title')} · 重点背诵", "", "> 保存从侧边栏整理的个人笔记；原始题目和参考答案仍留在本地重点数据中。"]
    progress_items = progress_payload.get("items") if isinstance(progress_payload.get("items"), dict) else {}
    for item_id, state in progress_items.items():
        record = items.get(str(item_id))
        if not record or record["subject"].get("id") != subject_id or not isinstance(state, dict):
            continue
        answer = str(state.get("answer") or "").strip()
        memory_note = str(state.get("memory_note") or "").strip()
        mastery = str(state.get("mastery") or "unseen")
        if not answer and not memory_note and mastery == "unseen":
            continue
        lines.extend(["", f"## {record.get('title')}", "", f"- 章节：{record['chapter'].get('title')}", f"- 题型：{record.get('type_label')}", f"- 掌握：{mastery}"])
        if answer:
            lines.extend(["", "### 我的作答", "", answer])
        if memory_note:
            lines.extend(["", "### 学习笔记", "", memory_note])
    target, storage, uri = oral_focus_notes_target(subject)
    atomic_write(target, "\n".join(lines).strip() + "\n")
    return target, storage, uri


def save_oral_focus_progress(item_id: str, answer: str, memory_note: str, mastery: str) -> dict:
    item = oral_focus_item_payload(item_id)
    if mastery not in {"unseen", "learning", "fuzzy", "mastered"}:
        raise ValueError("invalid oral focus mastery")
    answer = str(answer or "").replace("\r\n", "\n").strip()
    memory_note = str(memory_note or "").replace("\r\n", "\n").strip()
    if len(answer) > 50000 or len(memory_note) > 30000:
        raise ValueError("oral focus response is too long")
    with ORAL_FOCUS_LOCK:
        payload = load_oral_focus_progress()
        payload.setdefault("items", {})[item_id] = {
            "answer": answer,
            "memory_note": memory_note,
            "mastery": mastery,
            "updated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        }
        atomic_write(ORAL_FOCUS_PROGRESS_PATH, json.dumps(payload, ensure_ascii=False, indent=2))
        target, storage, uri = write_oral_focus_notes(item["subject"]["id"], payload)
    return {"saved": bool(answer or memory_note or mastery != "unseen"), "progress": payload["items"][item_id], "path": str(target), "storage": storage, "obsidian_uri": uri}


